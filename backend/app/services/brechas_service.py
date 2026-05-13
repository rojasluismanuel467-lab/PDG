from __future__ import annotations

import uuid
from dataclasses import dataclass
from datetime import UTC, datetime
from io import BytesIO

from sqlalchemy.orm import Session

from app.core.enums import PermissionLevel, ProjectArtifactStatus
from app.exceptions.domain import ForbiddenDomainError, NotFoundDomainError, ValidationDomainError
from app.models.gap_analysis_report import GapAnalysisReport, GapAnalysisReportVersion
from app.models.gaps_crud_matrix import GapsCRUDMatrix, GapsCRUDMatrixVersion
from app.models.integration_quality_rules import (
    IntegrationQualityRules,
    IntegrationQualityRulesVersion,
)
from app.models.project_artifact import ProjectArtifact
from app.repositories.brechas_repository import BrechasRepository
from app.schemas.brechas import (
    BrechaImpacto,
    BrechaPrioridad,
    BrechasVersionResponse,
    CRUDComparisonSchema,
    CRUDMatrixResponse,
    CRUDMatrixSnapshotRequest,
    ExportFormat,
    GapAnalysisReportResponse,
    GapAnalysisReportSnapshotRequest,
    GapSchema,
    IntegrationQualityRulesResponse,
    IntegrationQualityRulesSnapshotRequest,
    IntegrationRulePriority,
    IntegrationRuleSchema,
    IntegrationRuleType,
)
from app.services.project_permission_service import ProjectPermissionService


@dataclass
class _EntitySnapshot:
    nombre: str
    atributos: int
    fks: int


class BrechasService:
    CRUD_ARTIFACT_CODE = "GAPS_CRUD_MATRIX"
    GAP_REPORT_ARTIFACT_CODE = "GAPS_ANALYSIS_REPORT"
    RULES_ARTIFACT_CODE = "GAPS_INTEGRATION_QUALITY_RULES"
    SUPPORTED_CODES = {CRUD_ARTIFACT_CODE, GAP_REPORT_ARTIFACT_CODE, RULES_ARTIFACT_CODE}

    @staticmethod
    def _normalize(value: str) -> str:
        normalized = value.lower()
        return "".join(ch for ch in normalized if ch.isalnum())

    @staticmethod
    def _version_label(version_number: int) -> str:
        return f"{version_number}.0"

    @staticmethod
    def _artifact_lock_check(artifact: ProjectArtifact) -> None:
        if artifact.status == ProjectArtifactStatus.APPROVED:
            raise ForbiddenDomainError("Artifact is approved and locked for edits")

    @staticmethod
    def _map_versions(versions: list) -> list[BrechasVersionResponse]:
        ordered = sorted(versions, key=lambda item: item.version_number, reverse=True)
        return [
            BrechasVersionResponse(
                version=BrechasService._version_label(version.version_number),
                fecha=version.created_at,
                autor=version.created_by_user_email,
                descripcion_cambio=version.change_summary,
            )
            for version in ordered
        ]

    @classmethod
    def _resolve_artifact(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        expected_code: str,
    ) -> ProjectArtifact:
        artifact = BrechasRepository.get_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if artifact is None:
            raise NotFoundDomainError("Project artifact not found")
        if artifact.code not in cls.SUPPORTED_CODES:
            raise ValidationDomainError("This endpoint only supports Brechas artifacts")
        if artifact.code != expected_code:
            raise ValidationDomainError(f"Artifact code mismatch. Expected {expected_code}")
        return artifact

    @staticmethod
    def _resolve_permission(
        db: Session,
        *,
        project_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        artifact: ProjectArtifact,
        minimum_level: PermissionLevel,
    ) -> int:
        _, _, level = ProjectPermissionService.resolve_artifact_level(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=minimum_level,
        )
        return level

    @staticmethod
    def _fallback_entities(
        project_name: str,
    ) -> tuple[list[_EntitySnapshot], list[_EntitySnapshot]]:
        normalized = BrechasService._normalize(project_name)
        if "banco" in normalized or "fin" in normalized:
            asis = ["Cliente", "Cuenta", "Transaccion", "Producto", "ClienteProducto"]
            tobe = asis + ["Segmento", "Canal", "EventoDigital"]
        elif "energia" in normalized:
            asis = ["Activo", "Medidor", "OrdenTrabajo", "LecturaConsumo", "ClienteServicio"]
            tobe = asis + ["EventoSCADA", "CanalAtencion"]
        else:
            asis = ["Cliente", "Producto", "Orden", "Factura", "Pago"]
            tobe = asis + ["Canal", "Evento", "Segmento"]
        return (
            [_EntitySnapshot(nombre=name, atributos=6, fks=1) for name in asis],
            [_EntitySnapshot(nombre=name, atributos=8, fks=2) for name in tobe],
        )

    @staticmethod
    def _entity_snapshots_from_conceptual(model) -> list[_EntitySnapshot]:
        snapshots: list[_EntitySnapshot] = []
        for entity in model.entities:
            attrs = len(entity.attributes)
            fks = sum(1 for attribute in entity.attributes if attribute.is_fk)
            snapshots.append(_EntitySnapshot(nombre=entity.name, atributos=attrs, fks=fks))
        return snapshots

    @staticmethod
    def _entity_snapshots_from_logical(model) -> list[_EntitySnapshot]:
        snapshots: list[_EntitySnapshot] = []
        for table in model.tables:
            columns = table.get("columnas", [])
            attrs = len(columns)
            fks = sum(
                1 for column in columns if bool(column.get("es_fk")) or bool(column.get("is_fk"))
            )
            snapshots.append(
                _EntitySnapshot(
                    nombre=str(table.get("entidad_origen") or table.get("nombre") or "Entidad"),
                    atributos=attrs,
                    fks=fks,
                )
            )
        return snapshots

    @classmethod
    def _build_context(cls, db: Session, *, project_id: uuid.UUID) -> dict:
        project = BrechasRepository.get_project(db, project_id=project_id)
        if project is None:
            raise NotFoundDomainError("Project not found")

        artifacts = BrechasRepository.list_project_artifacts(db, project_id=project_id)
        artifact_by_code = {artifact.code: artifact for artifact in artifacts}

        asis_entities: list[_EntitySnapshot] = []
        tobe_entities: list[_EntitySnapshot] = []
        dfd_asis_flows = 0
        dfd_tobe_flows = 0

        asis_conceptual_artifact = artifact_by_code.get("ASIS_CONCEPTUAL_DIAGRAM")
        if asis_conceptual_artifact:
            model = BrechasRepository.get_conceptual_model_by_artifact(
                db,
                artifact_id=asis_conceptual_artifact.id,
            )
            if model:
                asis_entities = cls._entity_snapshots_from_conceptual(model)

        tobe_conceptual_artifact = artifact_by_code.get("TOBE_CONCEPTUAL_DIAGRAM")
        if tobe_conceptual_artifact:
            model = BrechasRepository.get_conceptual_model_by_artifact(
                db,
                artifact_id=tobe_conceptual_artifact.id,
            )
            if model:
                tobe_entities.extend(cls._entity_snapshots_from_conceptual(model))

        logical_artifact = artifact_by_code.get("TOBE_LOGICAL_DATA_MODEL")
        if logical_artifact:
            logical_model = BrechasRepository.get_logical_model_by_artifact(
                db,
                artifact_id=logical_artifact.id,
            )
            if logical_model:
                tobe_entities.extend(cls._entity_snapshots_from_logical(logical_model))

        dfd_asis_artifact = artifact_by_code.get("ASIS_DFD")
        if dfd_asis_artifact:
            dfd_model = BrechasRepository.get_dfd_model_by_artifact(
                db,
                artifact_id=dfd_asis_artifact.id,
            )
            if dfd_model:
                dfd_asis_flows = len(dfd_model.flows or [])

        dfd_tobe_artifact = artifact_by_code.get("TOBE_DFD")
        if dfd_tobe_artifact:
            dfd_model = BrechasRepository.get_dfd_model_by_artifact(
                db,
                artifact_id=dfd_tobe_artifact.id,
            )
            if dfd_model:
                dfd_tobe_flows = len(dfd_model.flows or [])

        if not asis_entities and not tobe_entities:
            fallback_asis, fallback_tobe = cls._fallback_entities(project.nombre)
            asis_entities = fallback_asis
            tobe_entities = fallback_tobe

        tobe_dedup: dict[str, _EntitySnapshot] = {}
        for snapshot in tobe_entities:
            key = cls._normalize(snapshot.nombre)
            if key and key not in tobe_dedup:
                tobe_dedup[key] = snapshot

        return {
            "project_id": project_id,
            "project_name": project.nombre,
            "asis": asis_entities,
            "tobe": list(tobe_dedup.values()),
            "dfd_asis_flows": dfd_asis_flows,
            "dfd_tobe_flows": dfd_tobe_flows,
        }

    @classmethod
    def _infer_ops(cls, snapshot: _EntitySnapshot, stage: str) -> dict[str, bool]:
        key = cls._normalize(snapshot.nombre)
        is_catalog = any(
            word in key for word in ("catalog", "categoria", "tipo", "segment", "canal", "param")
        )
        is_event = any(word in key for word in ("evento", "log", "audit", "hist", "transaccion"))
        is_bridge = snapshot.fks >= 2 and snapshot.atributos <= 7

        if stage == "asis":
            if is_catalog:
                return {"create": False, "read": True, "update": False, "delete": False}
            if is_event or is_bridge:
                return {"create": True, "read": True, "update": False, "delete": False}
            return {"create": True, "read": True, "update": True, "delete": False}

        if is_catalog:
            return {"create": True, "read": True, "update": True, "delete": False}
        if is_event or is_bridge:
            return {"create": True, "read": True, "update": True, "delete": False}
        return {"create": True, "read": True, "update": True, "delete": True}

    @classmethod
    def _crud_gap_description(
        cls,
        *,
        entidad: str,
        asis: _EntitySnapshot | None,
        tobe: _EntitySnapshot | None,
    ) -> str:
        if asis is None and tobe is not None:
            return (
                f"Nueva entidad en TO-BE: {entidad}. "
                "Se incorpora para cubrir trazabilidad y gobierno de datos."
            )
        if asis is not None and tobe is None:
            return (
                f"Entidad {entidad} presente en AS-IS y ausente en TO-BE. "
                "Se recomienda consolidacion o migracion funcional."
            )
        if asis is None or tobe is None:
            return "Brecha por informacion estructural incompleta."

        delta_attrs = tobe.atributos - asis.atributos
        delta_fks = tobe.fks - asis.fks
        parts: list[str] = []
        if delta_attrs > 0:
            parts.append(f"se agregan {delta_attrs} atributos")
        elif delta_attrs < 0:
            parts.append(f"se simplifican {abs(delta_attrs)} atributos")
        if delta_fks > 0:
            parts.append(f"se incrementa integracion con {delta_fks} referencias")
        elif delta_fks < 0:
            parts.append(f"se reducen {abs(delta_fks)} dependencias")

        if not parts:
            return f"No se observan cambios estructurales mayores en {entidad}."
        return f"Para {entidad}, " + " y ".join(parts) + "."

    @classmethod
    def _crud_impact(cls, row: CRUDComparisonSchema, *, is_critical_entity: bool) -> BrechaImpacto:
        score = 0
        if row.asis_create != row.tobe_create:
            score += 2
        if row.asis_read != row.tobe_read:
            score += 3
        if row.asis_update != row.tobe_update:
            score += 2
        if row.asis_delete != row.tobe_delete:
            score += 4
        if is_critical_entity:
            score += 2
        if not row.tobe_read:
            score += 2
        if not row.asis_delete and row.tobe_delete:
            score += 2
        if score >= 9:
            return BrechaImpacto.ALTO
        if score >= 5:
            return BrechaImpacto.MEDIO
        return BrechaImpacto.BAJO

    @classmethod
    def _generate_crud_matrix_payload(cls, db: Session, *, project_id: uuid.UUID) -> dict:
        context = cls._build_context(db, project_id=project_id)
        asis_map = {cls._normalize(item.nombre): item for item in context["asis"]}
        tobe_map = {cls._normalize(item.nombre): item for item in context["tobe"]}

        ordered_keys: list[str] = []
        for entity in context["asis"]:
            key = cls._normalize(entity.nombre)
            if key and key not in ordered_keys:
                ordered_keys.append(key)
        for entity in context["tobe"]:
            key = cls._normalize(entity.nombre)
            if key and key not in ordered_keys:
                ordered_keys.append(key)

        comparisons: list[dict] = []
        for index, key in enumerate(ordered_keys):
            asis = asis_map.get(key)
            tobe = tobe_map.get(key)
            entidad = tobe.nombre if tobe else asis.nombre if asis else f"Entidad {index + 1}"
            asis_ops = (
                cls._infer_ops(asis, "asis")
                if asis
                else {
                    "create": False,
                    "read": False,
                    "update": False,
                    "delete": False,
                }
            )
            tobe_ops = (
                cls._infer_ops(tobe, "tobe")
                if tobe
                else {
                    "create": False,
                    "read": False,
                    "update": False,
                    "delete": False,
                }
            )

            key_norm = cls._normalize(entidad)
            is_critical = any(
                token in key_norm
                for token in ("cliente", "cuenta", "transaccion", "pago", "factura", "riesgo")
            )
            provisional = CRUDComparisonSchema(
                id=f"crud-{index + 1}",
                entidad=entidad,
                asis_create=asis_ops["create"],
                asis_read=asis_ops["read"],
                asis_update=asis_ops["update"],
                asis_delete=asis_ops["delete"],
                tobe_create=tobe_ops["create"],
                tobe_read=tobe_ops["read"],
                tobe_update=tobe_ops["update"],
                tobe_delete=tobe_ops["delete"],
                brecha=cls._crud_gap_description(entidad=entidad, asis=asis, tobe=tobe),
                impacto=BrechaImpacto.BAJO,
            )
            impact = cls._crud_impact(provisional, is_critical_entity=is_critical)
            comparisons.append({**provisional.model_dump(mode="json"), "impacto": impact.value})

        description = (
            "Matriz comparativa CRUD contextualizada con artefactos AS-IS/TO-BE. "
            f"Flujos DFD detectados: AS-IS={context['dfd_asis_flows']} y TO-BE={context['dfd_tobe_flows']}."
        )
        return {
            "nombre": "Matriz CRUD Comparativa",
            "descripcion": description,
            "comparaciones": comparisons,
        }

    @classmethod
    def _build_gap_area(cls, entity_name: str) -> str:
        key = cls._normalize(entity_name)
        if any(token in key for token in ("segment", "catalog", "canal", "glosario")):
            return "Gobierno y Catalogo de Datos"
        if any(token in key for token in ("evento", "flujo", "integr", "api")):
            return "Integracion y Flujos de Datos"
        if any(token in key for token in ("calidad", "valid", "regla")):
            return "Calidad y Operacion de Datos"
        return "Modelo de Datos y Dominio"

    @classmethod
    def _build_gap_priority(cls, impact: BrechaImpacto, entity_name: str) -> BrechaPrioridad:
        key = cls._normalize(entity_name)
        is_critical_entity = any(
            token in key
            for token in ("cliente", "cuenta", "transaccion", "pago", "factura", "riesgo")
        )
        if impact == BrechaImpacto.ALTO and is_critical_entity:
            return BrechaPrioridad.CRITICA
        if impact == BrechaImpacto.ALTO:
            return BrechaPrioridad.ALTA
        if impact == BrechaImpacto.MEDIO:
            return BrechaPrioridad.MEDIA
        return BrechaPrioridad.BAJA

    @classmethod
    def _generate_gap_report_payload(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        crud_comparisons: list[dict] | None = None,
    ) -> dict:
        context = cls._build_context(db, project_id=project_id)
        rows = crud_comparisons
        if rows is None:
            rows = cls._generate_crud_matrix_payload(db, project_id=project_id)["comparaciones"]

        gaps: list[dict] = []
        for index, row in enumerate(rows):
            change_count = sum(
                int(bool(row.get(flag_left)) != bool(row.get(flag_right)))
                for flag_left, flag_right in (
                    ("asis_create", "tobe_create"),
                    ("asis_read", "tobe_read"),
                    ("asis_update", "tobe_update"),
                    ("asis_delete", "tobe_delete"),
                )
            )
            if change_count == 0 and not row.get("brecha"):
                continue
            impact = BrechaImpacto(str(row.get("impacto") or BrechaImpacto.BAJO.value))
            priority = cls._build_gap_priority(impact, str(row.get("entidad", "")))
            recommendation = (
                f"Definir plan de remediacion para {row.get('entidad', 'Entidad')} "
                f"con foco en controles y seguimiento para impacto {impact.value.lower()}."
            )
            gaps.append(
                GapSchema(
                    id=f"gap-{index + 1}",
                    area=cls._build_gap_area(str(row.get("entidad", ""))),
                    brecha=str(row.get("brecha") or "Brecha detectada en comparacion CRUD."),
                    impacto=impact,
                    prioridad=priority,
                    recomendacion=recommendation,
                ).model_dump(mode="json")
            )

        total_gaps = len(gaps)
        critical_gaps = sum(
            1 for item in gaps if item["prioridad"] == BrechaPrioridad.CRITICA.value
        )
        recommendations: list[str] = []
        for item in gaps:
            if item["recomendacion"] not in recommendations:
                recommendations.append(item["recomendacion"])

        summary = (
            f"El reporte consolida {total_gaps} brechas para {context['project_name']}. "
            f"Se identifican {critical_gaps} brechas criticas con impacto en gobierno, integracion y calidad."
        )

        return {
            "nombre": "Reporte de Analisis de Brechas",
            "descripcion": "Documento ejecutivo con brechas, impacto, prioridad y recomendaciones.",
            "resumen_ejecutivo": summary,
            "brechas": gaps,
            "total_brechas": total_gaps,
            "brechas_criticas": critical_gaps,
            "recomendaciones_prioritarias": recommendations[:5],
            "formato_objetivo": ["PDF", "WORD", "MARKDOWN"],
        }

    @classmethod
    def _rule_type_from_gap(cls, gap: dict) -> IntegrationRuleType:
        area = cls._normalize(str(gap.get("area", "")))
        brecha = cls._normalize(str(gap.get("brecha", "")))
        if "integr" in area or "flujo" in area:
            return IntegrationRuleType.CONSOLIDACION
        if "catalog" in area or "gobierno" in area or "match" in brecha:
            return IntegrationRuleType.MATCHING
        return IntegrationRuleType.VALIDACION

    @classmethod
    def _rule_priority_from_gap(cls, gap: dict) -> IntegrationRulePriority:
        priority = str(gap.get("prioridad", "Media"))
        if priority in (BrechaPrioridad.CRITICA.value, BrechaPrioridad.ALTA.value):
            return IntegrationRulePriority.ALTA
        if priority == BrechaPrioridad.MEDIA.value:
            return IntegrationRulePriority.MEDIA
        return IntegrationRulePriority.BAJA

    @classmethod
    def _generate_integration_rules_payload(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        gap_report: dict | None = None,
    ) -> dict:
        context = cls._build_context(db, project_id=project_id)
        report = gap_report or cls._generate_gap_report_payload(db, project_id=project_id)
        gaps = report.get("brechas", [])

        rules: list[dict] = []
        for index, gap in enumerate(gaps[:12]):
            rule_type = cls._rule_type_from_gap(gap)
            rule_priority = cls._rule_priority_from_gap(gap)
            area = str(gap.get("area", "Area"))
            gap_text = str(gap.get("brecha", "Brecha"))
            rule_name = f"Regla {rule_type.value} - {area}"
            condition = f"Si el registro pertenece al dominio {area} y presenta patron asociado a '{gap_text[:80]}'."
            action = (
                "Aplicar estandar de transformacion, validar campos obligatorios y "
                "registrar evidencia de cumplimiento en bitacora de calidad."
            )
            rules.append(
                IntegrationRuleSchema(
                    id=f"rule-{index + 1}",
                    nombre=rule_name,
                    descripcion=f"Regla derivada de brecha priorizada en {area}.",
                    tipo=rule_type,
                    prioridad=rule_priority,
                    condicion=condition,
                    accion=action,
                ).model_dump(mode="json")
            )

        acceptance_criteria = [
            "Cobertura de validaciones mayor o igual al 95% en dominios priorizados.",
            "Trazabilidad completa de reglas aplicadas por lote y por entidad.",
            "Alertas automaticas cuando la tasa de error supere el umbral definido.",
            "Evidencia de pruebas de integracion para cada regla de prioridad alta.",
        ]

        summary = (
            f"Documento tecnico para {context['project_name']} con reglas de matching, "
            "validacion y consolidacion orientadas a cerrar brechas de arquitectura de datos."
        )

        return {
            "nombre": "Reglas de Integracion y Calidad de Datos",
            "descripcion": "Especificacion tecnica para consolidar datos y asegurar integridad.",
            "resumen_tecnico": summary,
            "reglas": rules,
            "criterios_aceptacion": acceptance_criteria,
            "formato_objetivo": ["MARKDOWN", "WORD", "PDF"],
        }

    @classmethod
    def _ensure_crud_matrix(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> GapsCRUDMatrix:
        model = BrechasRepository.get_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if model is not None:
            return model

        payload = cls._generate_crud_matrix_payload(db, project_id=project_id)
        now = datetime.now(UTC)
        model = BrechasRepository.create_crud_matrix(
            db,
            model=GapsCRUDMatrix(
                project_id=project_id,
                artifact_id=artifact_id,
                name=payload["nombre"],
                description=payload["descripcion"],
                comparisons=payload["comparaciones"],
                current_version_number=1,
                created_by_user_id=actor_user_id,
                updated_by_user_id=actor_user_id,
                last_saved_at=now,
            ),
        )
        BrechasRepository.create_crud_matrix_version(
            db,
            version=GapsCRUDMatrixVersion(
                matrix_id=model.id,
                version_number=1,
                snapshot=payload,
                change_summary="Initial CRUD matrix",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.flush()
        refreshed = BrechasRepository.get_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        return refreshed or model

    @classmethod
    def _ensure_gap_report(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> GapAnalysisReport:
        model = BrechasRepository.get_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if model is not None:
            return model

        payload = cls._generate_gap_report_payload(db, project_id=project_id)
        now = datetime.now(UTC)
        model = BrechasRepository.create_gap_report(
            db,
            model=GapAnalysisReport(
                project_id=project_id,
                artifact_id=artifact_id,
                name=payload["nombre"],
                description=payload["descripcion"],
                executive_summary=payload["resumen_ejecutivo"],
                gaps=payload["brechas"],
                total_gaps=payload["total_brechas"],
                critical_gaps=payload["brechas_criticas"],
                priority_recommendations=payload["recomendaciones_prioritarias"],
                target_formats=payload["formato_objetivo"],
                current_version_number=1,
                created_by_user_id=actor_user_id,
                updated_by_user_id=actor_user_id,
                last_saved_at=now,
            ),
        )
        BrechasRepository.create_gap_report_version(
            db,
            version=GapAnalysisReportVersion(
                report_id=model.id,
                version_number=1,
                snapshot=payload,
                change_summary="Initial gap analysis report",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.flush()
        refreshed = BrechasRepository.get_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        return refreshed or model

    @classmethod
    def _ensure_integration_rules(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> IntegrationQualityRules:
        model = BrechasRepository.get_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if model is not None:
            return model

        payload = cls._generate_integration_rules_payload(db, project_id=project_id)
        now = datetime.now(UTC)
        model = BrechasRepository.create_integration_rules(
            db,
            model=IntegrationQualityRules(
                project_id=project_id,
                artifact_id=artifact_id,
                name=payload["nombre"],
                description=payload["descripcion"],
                technical_summary=payload["resumen_tecnico"],
                rules=payload["reglas"],
                acceptance_criteria=payload["criterios_aceptacion"],
                target_formats=payload["formato_objetivo"],
                current_version_number=1,
                created_by_user_id=actor_user_id,
                updated_by_user_id=actor_user_id,
                last_saved_at=now,
            ),
        )
        BrechasRepository.create_integration_rules_version(
            db,
            version=IntegrationQualityRulesVersion(
                document_id=model.id,
                version_number=1,
                snapshot=payload,
                change_summary="Initial integration and quality rules document",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.flush()
        refreshed = BrechasRepository.get_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        return refreshed or model

    @classmethod
    def _map_crud_matrix(cls, model: GapsCRUDMatrix) -> CRUDMatrixResponse:
        return CRUDMatrixResponse(
            id=model.id,
            proyecto_id=model.project_id,
            entregable_id=model.artifact_id,
            nombre=model.name,
            descripcion=model.description,
            comparaciones=[
                CRUDComparisonSchema.model_validate(comparison) for comparison in model.comparisons
            ],
            version_actual=cls._version_label(model.current_version_number),
            historial_versiones=cls._map_versions(model.versions),
            created_at=model.created_at,
            updated_at=model.updated_at,
        )

    @classmethod
    def _map_gap_report(cls, model: GapAnalysisReport) -> GapAnalysisReportResponse:
        return GapAnalysisReportResponse(
            id=model.id,
            proyecto_id=model.project_id,
            entregable_id=model.artifact_id,
            nombre=model.name,
            descripcion=model.description,
            resumen_ejecutivo=model.executive_summary,
            brechas=[GapSchema.model_validate(gap) for gap in model.gaps],
            total_brechas=model.total_gaps,
            brechas_criticas=model.critical_gaps,
            recomendaciones_prioritarias=model.priority_recommendations,
            formato_objetivo=model.target_formats,
            version_actual=cls._version_label(model.current_version_number),
            historial_versiones=cls._map_versions(model.versions),
            created_at=model.created_at,
            updated_at=model.updated_at,
        )

    @classmethod
    def _map_integration_rules(
        cls,
        model: IntegrationQualityRules,
    ) -> IntegrationQualityRulesResponse:
        return IntegrationQualityRulesResponse(
            id=model.id,
            proyecto_id=model.project_id,
            entregable_id=model.artifact_id,
            nombre=model.name,
            descripcion=model.description,
            resumen_tecnico=model.technical_summary,
            reglas=[IntegrationRuleSchema.model_validate(rule) for rule in model.rules],
            criterios_aceptacion=model.acceptance_criteria,
            formato_objetivo=model.target_formats,
            version_actual=cls._version_label(model.current_version_number),
            historial_versiones=cls._map_versions(model.versions),
            created_at=model.created_at,
            updated_at=model.updated_at,
        )

    @classmethod
    def get_crud_matrix(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> CRUDMatrixResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.CRUD_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.LECTURA,
        )
        model = cls._ensure_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        db.commit()
        return cls._map_crud_matrix(model)

    @classmethod
    def upsert_crud_matrix(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
        payload: CRUDMatrixSnapshotRequest,
    ) -> CRUDMatrixResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.CRUD_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.EDITAR,
        )
        cls._artifact_lock_check(artifact)
        model = cls._ensure_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )

        model.name = payload.nombre
        model.description = payload.descripcion
        model.comparisons = [
            comparison.model_dump(mode="json") for comparison in payload.comparaciones
        ]
        model.current_version_number += 1
        model.updated_by_user_id = actor_user_id
        model.last_saved_at = datetime.now(UTC)

        BrechasRepository.create_crud_matrix_version(
            db,
            version=GapsCRUDMatrixVersion(
                matrix_id=model.id,
                version_number=model.current_version_number,
                snapshot=payload.model_dump(mode="json", exclude={"change_summary"}),
                change_summary=payload.change_summary or "Manual CRUD matrix update.",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.commit()
        refreshed = BrechasRepository.get_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if refreshed is None:
            raise NotFoundDomainError("CRUD matrix not found after save")
        return cls._map_crud_matrix(refreshed)

    @classmethod
    def generate_crud_matrix(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> CRUDMatrixResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.CRUD_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.EDITAR,
        )
        cls._artifact_lock_check(artifact)
        model = cls._ensure_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        generated = cls._generate_crud_matrix_payload(db, project_id=project_id)
        model.name = generated["nombre"]
        model.description = generated["descripcion"]
        model.comparisons = generated["comparaciones"]
        model.current_version_number += 1
        model.updated_by_user_id = actor_user_id
        model.last_saved_at = datetime.now(UTC)

        BrechasRepository.create_crud_matrix_version(
            db,
            version=GapsCRUDMatrixVersion(
                matrix_id=model.id,
                version_number=model.current_version_number,
                snapshot=generated,
                change_summary="Generated contextual CRUD matrix.",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.commit()
        refreshed = BrechasRepository.get_crud_matrix(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if refreshed is None:
            raise NotFoundDomainError("CRUD matrix not found after generation")
        return cls._map_crud_matrix(refreshed)

    @classmethod
    def get_gap_report(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> GapAnalysisReportResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.GAP_REPORT_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.LECTURA,
        )
        model = cls._ensure_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        db.commit()
        return cls._map_gap_report(model)

    @classmethod
    def upsert_gap_report(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
        payload: GapAnalysisReportSnapshotRequest,
    ) -> GapAnalysisReportResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.GAP_REPORT_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.EDITAR,
        )
        cls._artifact_lock_check(artifact)
        model = cls._ensure_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        model.name = payload.nombre
        model.description = payload.descripcion
        model.executive_summary = payload.resumen_ejecutivo
        model.gaps = [gap.model_dump(mode="json") for gap in payload.brechas]
        model.total_gaps = payload.total_brechas
        model.critical_gaps = payload.brechas_criticas
        model.priority_recommendations = payload.recomendaciones_prioritarias
        model.target_formats = payload.formato_objetivo
        model.current_version_number += 1
        model.updated_by_user_id = actor_user_id
        model.last_saved_at = datetime.now(UTC)

        BrechasRepository.create_gap_report_version(
            db,
            version=GapAnalysisReportVersion(
                report_id=model.id,
                version_number=model.current_version_number,
                snapshot=payload.model_dump(mode="json", exclude={"change_summary"}),
                change_summary=payload.change_summary or "Manual gap report update.",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.commit()
        refreshed = BrechasRepository.get_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if refreshed is None:
            raise NotFoundDomainError("Gap report not found after save")
        return cls._map_gap_report(refreshed)

    @classmethod
    def generate_gap_report(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> GapAnalysisReportResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.GAP_REPORT_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.EDITAR,
        )
        cls._artifact_lock_check(artifact)
        model = cls._ensure_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )

        generated = cls._generate_gap_report_payload(db, project_id=project_id)
        model.name = generated["nombre"]
        model.description = generated["descripcion"]
        model.executive_summary = generated["resumen_ejecutivo"]
        model.gaps = generated["brechas"]
        model.total_gaps = generated["total_brechas"]
        model.critical_gaps = generated["brechas_criticas"]
        model.priority_recommendations = generated["recomendaciones_prioritarias"]
        model.target_formats = generated["formato_objetivo"]
        model.current_version_number += 1
        model.updated_by_user_id = actor_user_id
        model.last_saved_at = datetime.now(UTC)

        BrechasRepository.create_gap_report_version(
            db,
            version=GapAnalysisReportVersion(
                report_id=model.id,
                version_number=model.current_version_number,
                snapshot=generated,
                change_summary="Generated contextual gap analysis report.",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.commit()
        refreshed = BrechasRepository.get_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if refreshed is None:
            raise NotFoundDomainError("Gap report not found after generation")
        return cls._map_gap_report(refreshed)

    @classmethod
    def get_integration_rules(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> IntegrationQualityRulesResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.RULES_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.LECTURA,
        )
        model = cls._ensure_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        db.commit()
        return cls._map_integration_rules(model)

    @classmethod
    def upsert_integration_rules(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
        payload: IntegrationQualityRulesSnapshotRequest,
    ) -> IntegrationQualityRulesResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.RULES_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.EDITAR,
        )
        cls._artifact_lock_check(artifact)
        model = cls._ensure_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        model.name = payload.nombre
        model.description = payload.descripcion
        model.technical_summary = payload.resumen_tecnico
        model.rules = [rule.model_dump(mode="json") for rule in payload.reglas]
        model.acceptance_criteria = payload.criterios_aceptacion
        model.target_formats = payload.formato_objetivo
        model.current_version_number += 1
        model.updated_by_user_id = actor_user_id
        model.last_saved_at = datetime.now(UTC)

        BrechasRepository.create_integration_rules_version(
            db,
            version=IntegrationQualityRulesVersion(
                document_id=model.id,
                version_number=model.current_version_number,
                snapshot=payload.model_dump(mode="json", exclude={"change_summary"}),
                change_summary=payload.change_summary or "Manual integration rules update.",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.commit()
        refreshed = BrechasRepository.get_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if refreshed is None:
            raise NotFoundDomainError("Integration rules document not found after save")
        return cls._map_integration_rules(refreshed)

    @classmethod
    def generate_integration_rules(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
    ) -> IntegrationQualityRulesResponse:
        artifact = cls._resolve_artifact(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            expected_code=cls.RULES_ARTIFACT_CODE,
        )
        cls._resolve_permission(
            db,
            project_id=project_id,
            actor_user_id=actor_user_id,
            artifact=artifact,
            minimum_level=PermissionLevel.EDITAR,
        )
        cls._artifact_lock_check(artifact)
        model = cls._ensure_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        generated = cls._generate_integration_rules_payload(db, project_id=project_id)
        model.name = generated["nombre"]
        model.description = generated["descripcion"]
        model.technical_summary = generated["resumen_tecnico"]
        model.rules = generated["reglas"]
        model.acceptance_criteria = generated["criterios_aceptacion"]
        model.target_formats = generated["formato_objetivo"]
        model.current_version_number += 1
        model.updated_by_user_id = actor_user_id
        model.last_saved_at = datetime.now(UTC)

        BrechasRepository.create_integration_rules_version(
            db,
            version=IntegrationQualityRulesVersion(
                document_id=model.id,
                version_number=model.current_version_number,
                snapshot=generated,
                change_summary="Generated contextual integration and quality rules.",
                created_by_user_id=actor_user_id,
                created_by_user_email=actor_user_email,
            ),
        )
        db.commit()
        refreshed = BrechasRepository.get_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
        )
        if refreshed is None:
            raise NotFoundDomainError("Integration rules document not found after generation")
        return cls._map_integration_rules(refreshed)

    @staticmethod
    def _gap_report_markdown(report: GapAnalysisReportResponse) -> str:
        lines: list[str] = [
            f"# {report.nombre}",
            "",
            f"Proyecto: {report.proyecto_id}",
            f"Version: {report.version_actual}",
            "",
            "## Resumen Ejecutivo",
            report.resumen_ejecutivo,
            "",
            "## Metricas",
            f"- Total de brechas: {report.total_brechas}",
            f"- Brechas criticas: {report.brechas_criticas}",
            "",
            "## Brechas Identificadas",
            "| Area | Brecha | Impacto | Prioridad | Recomendacion |",
            "| --- | --- | --- | --- | --- |",
        ]
        for gap in report.brechas:
            lines.append(
                f"| {gap.area} | {gap.brecha} | {gap.impacto.value} | {gap.prioridad.value} | {gap.recomendacion} |"
            )
        if report.recomendaciones_prioritarias:
            lines.extend(["", "## Recomendaciones Prioritarias"])
            for item in report.recomendaciones_prioritarias:
                lines.append(f"- {item}")
        return "\n".join(lines)

    @staticmethod
    def _integration_rules_markdown(document: IntegrationQualityRulesResponse) -> str:
        lines: list[str] = [
            f"# {document.nombre}",
            "",
            f"Proyecto: {document.proyecto_id}",
            f"Version: {document.version_actual}",
            "",
            "## Resumen Tecnico",
            document.resumen_tecnico,
            "",
            "## Reglas",
            "| Regla | Tipo | Prioridad | Condicion | Accion |",
            "| --- | --- | --- | --- | --- |",
        ]
        for rule in document.reglas:
            lines.append(
                f"| {rule.nombre} | {rule.tipo.value} | {rule.prioridad.value} | {rule.condicion} | {rule.accion} |"
            )
        if document.criterios_aceptacion:
            lines.extend(["", "## Criterios de Aceptacion"])
            for item in document.criterios_aceptacion:
                lines.append(f"- {item}")
        return "\n".join(lines)

    @staticmethod
    def _build_gap_report_pdf(report: GapAnalysisReportResponse) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise ValidationDomainError(
                "PDF export dependency is missing. Install reportlab in backend."
            ) from exc

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), leftMargin=24, rightMargin=24)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(report.nombre, styles["Title"]),
            Paragraph(
                f"Proyecto: {report.proyecto_id} - Version: {report.version_actual}",
                styles["BodyText"],
            ),
            Spacer(1, 8),
            Paragraph("Resumen Ejecutivo", styles["Heading2"]),
            Paragraph(report.resumen_ejecutivo, styles["BodyText"]),
            Spacer(1, 8),
            Paragraph("Brechas Identificadas", styles["Heading2"]),
        ]

        data = [
            [
                Paragraph("<b>Area</b>", styles["BodyText"]),
                Paragraph("<b>Brecha</b>", styles["BodyText"]),
                Paragraph("<b>Impacto</b>", styles["BodyText"]),
                Paragraph("<b>Prioridad</b>", styles["BodyText"]),
                Paragraph("<b>Recomendacion</b>", styles["BodyText"]),
            ]
        ]
        for gap in report.brechas:
            data.append(
                [
                    Paragraph(gap.area, styles["BodyText"]),
                    Paragraph(gap.brecha, styles["BodyText"]),
                    Paragraph(gap.impacto.value, styles["BodyText"]),
                    Paragraph(gap.prioridad.value, styles["BodyText"]),
                    Paragraph(gap.recomendacion, styles["BodyText"]),
                ]
            )
        table = Table(data, colWidths=[110, 210, 70, 70, 250], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e5f4f8")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#102a43")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 10))
        story.append(Paragraph("Recomendaciones Prioritarias", styles["Heading2"]))
        for item in report.recomendaciones_prioritarias:
            story.append(Paragraph(f"- {item}", styles["BodyText"]))

        doc.build(story)
        return buffer.getvalue()

    @staticmethod
    def _build_integration_rules_pdf(document: IntegrationQualityRulesResponse) -> bytes:
        try:
            from reportlab.lib import colors
            from reportlab.lib.pagesizes import A4, landscape
            from reportlab.lib.styles import getSampleStyleSheet
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
        except ImportError as exc:
            raise ValidationDomainError(
                "PDF export dependency is missing. Install reportlab in backend."
            ) from exc

        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=landscape(A4), leftMargin=24, rightMargin=24)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(document.nombre, styles["Title"]),
            Paragraph(
                f"Proyecto: {document.proyecto_id} - Version: {document.version_actual}",
                styles["BodyText"],
            ),
            Spacer(1, 8),
            Paragraph("Resumen Tecnico", styles["Heading2"]),
            Paragraph(document.resumen_tecnico, styles["BodyText"]),
            Spacer(1, 8),
            Paragraph("Reglas de Integracion", styles["Heading2"]),
        ]
        data = [
            [
                Paragraph("<b>Regla</b>", styles["BodyText"]),
                Paragraph("<b>Tipo</b>", styles["BodyText"]),
                Paragraph("<b>Prioridad</b>", styles["BodyText"]),
                Paragraph("<b>Condicion</b>", styles["BodyText"]),
                Paragraph("<b>Accion</b>", styles["BodyText"]),
            ]
        ]
        for rule in document.reglas:
            data.append(
                [
                    Paragraph(rule.nombre, styles["BodyText"]),
                    Paragraph(rule.tipo.value, styles["BodyText"]),
                    Paragraph(rule.prioridad.value, styles["BodyText"]),
                    Paragraph(rule.condicion, styles["BodyText"]),
                    Paragraph(rule.accion, styles["BodyText"]),
                ]
            )
        table = Table(data, colWidths=[170, 90, 80, 220, 210], repeatRows=1)
        table.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8f8ef")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#0f5132")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        story.append(table)
        story.append(Spacer(1, 10))
        story.append(Paragraph("Criterios de Aceptacion", styles["Heading2"]))
        for item in document.criterios_aceptacion:
            story.append(Paragraph(f"- {item}", styles["BodyText"]))
        doc.build(story)
        return buffer.getvalue()

    @staticmethod
    def _build_gap_report_docx(report: GapAnalysisReportResponse) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:
            raise ValidationDomainError(
                "Word export dependency is missing. Install python-docx in backend."
            ) from exc

        document = Document()
        document.add_heading(report.nombre, level=1)
        document.add_paragraph(f"Proyecto: {report.proyecto_id}")
        document.add_paragraph(f"Version: {report.version_actual}")
        document.add_heading("Resumen Ejecutivo", level=2)
        document.add_paragraph(report.resumen_ejecutivo)
        document.add_heading("Brechas Identificadas", level=2)
        table = document.add_table(rows=1, cols=5)
        table.style = "Light List Accent 1"
        headers = table.rows[0].cells
        headers[0].text = "Area"
        headers[1].text = "Brecha"
        headers[2].text = "Impacto"
        headers[3].text = "Prioridad"
        headers[4].text = "Recomendacion"
        for gap in report.brechas:
            row = table.add_row().cells
            row[0].text = gap.area
            row[1].text = gap.brecha
            row[2].text = gap.impacto.value
            row[3].text = gap.prioridad.value
            row[4].text = gap.recomendacion
        document.add_heading("Recomendaciones Prioritarias", level=2)
        for item in report.recomendaciones_prioritarias:
            document.add_paragraph(item, style="List Bullet")
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @staticmethod
    def _build_integration_rules_docx(document_data: IntegrationQualityRulesResponse) -> bytes:
        try:
            from docx import Document
        except ImportError as exc:
            raise ValidationDomainError(
                "Word export dependency is missing. Install python-docx in backend."
            ) from exc

        document = Document()
        document.add_heading(document_data.nombre, level=1)
        document.add_paragraph(f"Proyecto: {document_data.proyecto_id}")
        document.add_paragraph(f"Version: {document_data.version_actual}")
        document.add_heading("Resumen Tecnico", level=2)
        document.add_paragraph(document_data.resumen_tecnico)
        document.add_heading("Reglas", level=2)
        table = document.add_table(rows=1, cols=5)
        table.style = "Light List Accent 1"
        headers = table.rows[0].cells
        headers[0].text = "Regla"
        headers[1].text = "Tipo"
        headers[2].text = "Prioridad"
        headers[3].text = "Condicion"
        headers[4].text = "Accion"
        for rule in document_data.reglas:
            row = table.add_row().cells
            row[0].text = rule.nombre
            row[1].text = rule.tipo.value
            row[2].text = rule.prioridad.value
            row[3].text = rule.condicion
            row[4].text = rule.accion
        document.add_heading("Criterios de Aceptacion", level=2)
        for item in document_data.criterios_aceptacion:
            document.add_paragraph(item, style="List Bullet")
        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()

    @classmethod
    def export_gap_report(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
        export_format: ExportFormat,
    ) -> tuple[bytes, str, str]:
        report = cls.get_gap_report(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        version = report.version_actual.replace(".", "-")
        base_name = f"gap-analysis-report-{project_id}-v{version}"

        if export_format == ExportFormat.MARKDOWN:
            content = cls._gap_report_markdown(report).encode("utf-8")
            return content, "text/markdown; charset=utf-8", f"{base_name}.md"
        if export_format == ExportFormat.PDF:
            content = cls._build_gap_report_pdf(report)
            return content, "application/pdf", f"{base_name}.pdf"
        if export_format == ExportFormat.WORD:
            content = cls._build_gap_report_docx(report)
            return (
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{base_name}.docx",
            )
        raise ValidationDomainError("Unsupported export format")

    @classmethod
    def export_integration_rules(
        cls,
        db: Session,
        *,
        project_id: uuid.UUID,
        artifact_id: uuid.UUID,
        actor_user_id: uuid.UUID,
        actor_user_email: str,
        export_format: ExportFormat,
    ) -> tuple[bytes, str, str]:
        document = cls.get_integration_rules(
            db,
            project_id=project_id,
            artifact_id=artifact_id,
            actor_user_id=actor_user_id,
            actor_user_email=actor_user_email,
        )
        version = document.version_actual.replace(".", "-")
        base_name = f"integration-quality-rules-{project_id}-v{version}"

        if export_format == ExportFormat.MARKDOWN:
            content = cls._integration_rules_markdown(document).encode("utf-8")
            return content, "text/markdown; charset=utf-8", f"{base_name}.md"
        if export_format == ExportFormat.PDF:
            content = cls._build_integration_rules_pdf(document)
            return content, "application/pdf", f"{base_name}.pdf"
        if export_format == ExportFormat.WORD:
            content = cls._build_integration_rules_docx(document)
            return (
                content,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                f"{base_name}.docx",
            )
        raise ValidationDomainError("Unsupported export format")
